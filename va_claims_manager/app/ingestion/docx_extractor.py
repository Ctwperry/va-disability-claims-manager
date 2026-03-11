"""
Word document (.docx) text extraction.
"""
import logging
from pathlib import Path

log = logging.getLogger(__name__)

# Approximate characters per "virtual page" when chunking Word docs
CHARS_PER_PAGE = 3000


def extract_pages(filepath: str | Path) -> list[dict]:
    """
    Extract text from a .docx file.
    Word documents have no real pages, so we chunk the text into
    virtual pages of ~3000 characters each (roughly one printed page).

    Returns: [{'page_number': n, 'raw_text': '...', 'has_image': False}, ...]
    """
    from docx import Document

    try:
        doc = Document(str(filepath))
    except Exception as exc:
        log.error("Cannot open DOCX %s: %s", filepath, exc)
        raise

    # Gather all paragraph text
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs)

    if not full_text.strip():
        return [{"page_number": 1, "raw_text": "", "has_image": False}]

    # Chunk into virtual pages
    pages = []
    chunks = _chunk_text(full_text, CHARS_PER_PAGE)
    for i, chunk in enumerate(chunks, start=1):
        pages.append({"page_number": i, "raw_text": chunk, "has_image": False})

    return pages


def _chunk_text(text: str, chunk_size: int) -> list[str]:
    """Split text into chunks of roughly chunk_size characters, breaking on newlines."""
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end >= len(text):
            chunks.append(text[start:])
            break
        # Try to break on a newline near the chunk boundary
        nl = text.rfind("\n", start, end)
        if nl > start:
            chunks.append(text[start:nl])
            start = nl + 1
        else:
            chunks.append(text[start:end])
            start = end
    return chunks
